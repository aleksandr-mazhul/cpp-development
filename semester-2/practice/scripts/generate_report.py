#!/usr/bin/env python3
"""Generate GOST-styled DOCX report for labs 1-4."""

from __future__ import annotations

import re
from pathlib import Path
from typing import List

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
ASSETS = ROOT / "tmp" / "report_assets"
OUTPUT = ROOT / "Отчет_практика_лабы_1-4.docx"

MAIN_FONT = "Times New Roman"
CODE_FONT = "Consolas"

FIG_COUNTER = 0
TAB_COUNTER = 0


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _shd(color_hex: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    return shd


def _border(size: int = 6, color: str = "999999"):
    pbdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), color)
        pbdr.append(el)
    return pbdr


def _apply_run_font(run, name: str, size_pt: int, bold: bool = False, italic: bool = False, color: str | None = None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rFonts.set(qn("w:eastAsia"), name)


# ---------------------------------------------------------------------------
# Page and style setup
# ---------------------------------------------------------------------------

def setup_sections(doc: Document):
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)

        sect_pr = section._sectPr
        title_pg = OxmlElement("w:titlePg")
        sect_pr.append(title_pg)

        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run()
        _apply_run_font(run, MAIN_FONT, 12)
        for tag in ("begin", "separate", "end"):
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), tag)
            if tag == "begin":
                run._r.append(fld)
                instr = OxmlElement("w:instrText")
                instr.set(qn("xml:space"), "preserve")
                instr.text = "PAGE"
                run._r.append(instr)
            else:
                run._r.append(fld)


def setup_styles(doc: Document):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = MAIN_FONT
    normal.font.size = Pt(14)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    h1 = styles["Heading 1"]
    h1.font.name = MAIN_FONT
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    hp1 = h1.paragraph_format
    hp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp1.first_line_indent = Cm(0)
    hp1.space_before = Pt(0)
    hp1.space_after = Pt(12)
    hp1.page_break_before = True
    hp1.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    h2 = styles["Heading 2"]
    h2.font.name = MAIN_FONT
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    hp2 = h2.paragraph_format
    hp2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp2.first_line_indent = Cm(0)
    hp2.space_before = Pt(12)
    hp2.space_after = Pt(6)
    hp2.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    hp2.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = MAIN_FONT
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.italic = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    hp3 = h3.paragraph_format
    hp3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp3.first_line_indent = Cm(0)
    hp3.space_before = Pt(6)
    hp3.space_after = Pt(3)
    hp3.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    hp3.keep_with_next = True

    code_style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = CODE_FONT
    code_style.font.size = Pt(10)
    code_style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    cpf = code_style.paragraph_format
    cpf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cpf.first_line_indent = Cm(0)
    cpf.left_indent = Cm(0.3)
    cpf.right_indent = Cm(0)
    cpf.space_before = Pt(3)
    cpf.space_after = Pt(3)
    cpf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pPr = code_style.element.get_or_add_pPr()
    pPr.append(_shd("F5F5F5"))
    pPr.append(_border())

    cap_style = styles.add_style("FigureCaption", WD_STYLE_TYPE.PARAGRAPH)
    cap_style.font.name = MAIN_FONT
    cap_style.font.size = Pt(12)
    cap_style.font.italic = True
    cpf2 = cap_style.paragraph_format
    cpf2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cpf2.first_line_indent = Cm(0)
    cpf2.space_before = Pt(3)
    cpf2.space_after = Pt(9)
    cpf2.keep_with_next = False

    tcap_style = styles.add_style("TableCaption", WD_STYLE_TYPE.PARAGRAPH)
    tcap_style.font.name = MAIN_FONT
    tcap_style.font.size = Pt(12)
    tcap_style.font.italic = True
    tp = tcap_style.paragraph_format
    tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tp.first_line_indent = Cm(0)
    tp.space_before = Pt(6)
    tp.space_after = Pt(3)
    tp.keep_with_next = True


def enable_update_fields(doc: Document):
    settings = doc.settings.element
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)


# ---------------------------------------------------------------------------
# Content builders
# ---------------------------------------------------------------------------

def add_para(doc: Document, text: str, indent: bool = True, bold: bool = False, alignment=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Cm(1.25) if indent else Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.alignment = alignment or WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    _apply_run_font(run, MAIN_FONT, 14, bold=bold)
    return p


def add_bullets(doc: Document, items: List[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.left_indent = Cm(1.25)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in p.runs:
            _apply_run_font(r, MAIN_FONT, 14)
        run = p.add_run(item) if not p.runs else p.runs[0]
        if not p.runs or run.text == "":
            run = p.add_run(item)
        else:
            run.text = item
        _apply_run_font(run, MAIN_FONT, 14)


def add_code(doc: Document, code: str):
    text = code.replace("\t", "    ").rstrip()
    lines = text.splitlines() or [""]
    for line in lines:
        p = doc.add_paragraph(style="CodeBlock")
        run = p.add_run(line if line else " ")
        _apply_run_font(run, CODE_FONT, 10)


def add_figure(doc: Document, image_path: Path, caption: str, width_inches: float = 5.5):
    global FIG_COUNTER
    FIG_COUNTER += 1
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(width_inches))
    cap = doc.add_paragraph(style="FigureCaption")
    run = cap.add_run(f"Рисунок {FIG_COUNTER} — {caption}")
    _apply_run_font(run, MAIN_FONT, 12, italic=True)


def add_table_caption(doc: Document, caption: str):
    global TAB_COUNTER
    TAB_COUNTER += 1
    p = doc.add_paragraph(style="TableCaption")
    run = p.add_run(f"Таблица {TAB_COUNTER} — {caption}")
    _apply_run_font(run, MAIN_FONT, 12, italic=True)


def add_table(doc: Document, headers: List[str], rows: List[List[str]], caption: str,
              col_widths: List[float] | None = None):
    add_table_caption(doc, caption)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(h)
        _apply_run_font(run, MAIN_FONT, 12, bold=True)
        tcPr = cell._tc.get_or_add_tcPr()
        tcPr.append(_shd("E8E8E8"))
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(val)
            _apply_run_font(run, MAIN_FONT, 12)

    if col_widths:
        for row in table.rows:
            for cell, w in zip(row.cells, col_widths):
                cell.width = Cm(w)
    empty = doc.add_paragraph()
    empty.paragraph_format.space_after = Pt(6)


# ---------------------------------------------------------------------------
# Title page and TOC
# ---------------------------------------------------------------------------

def add_title_page(doc: Document):
    def center(text: str, size=14, bold=False, space_after=0, italic=False):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Cm(0)
        pf.space_after = Pt(space_after)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(text)
        _apply_run_font(run, MAIN_FONT, size, bold=bold, italic=italic)
        return p

    def blank(space_after=12):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.first_line_indent = Cm(0)

    center("МИНИСТЕРСТВО ОБРАЗОВАНИЯ РЕСПУБЛИКИ БЕЛАРУСЬ", 14, bold=False, space_after=0)
    center("БЕЛОРУССКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ", 14, bold=True, space_after=0)
    center("МЕХАНИКО-МАТЕМАТИЧЕСКИЙ ФАКУЛЬТЕТ", 14, bold=False, space_after=0)
    center("Кафедра веб-технологий и компьютерного моделирования", 14, bold=False, space_after=0)

    blank(120)
    center("ВЫЧИСЛИТЕЛЬНАЯ ПРАКТИКА", 16, bold=True, space_after=24)
    center("ОТЧЁТ", 20, bold=True, space_after=6)

    blank(80)
    center("[ФИО СТУДЕНТА]", 14, bold=False, space_after=0)
    center("студента 2-го курса,", 14, bold=False, space_after=0)
    center("специальность «Математика и компьютерные науки»", 14, bold=False, space_after=0)

    blank(48)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run("Руководитель:\nстарший преподаватель Аленский Н. А.")
    _apply_run_font(run, MAIN_FONT, 14)

    blank(80)
    center("Минск, 2026", 14, bold=False, space_after=0)


def add_toc(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.page_break_before = True
    run = p.add_run("СОДЕРЖАНИЕ")
    _apply_run_font(run, MAIN_FONT, 16, bold=True)

    p2 = doc.add_paragraph()
    p2.paragraph_format.first_line_indent = Cm(0)
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p2.add_run()
    _apply_run_font(run, MAIN_FONT, 14)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    hint = OxmlElement("w:t")
    hint.text = "Обновите поле оглавления (F9)"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(hint)
    run._r.append(fld_end)


# ---------------------------------------------------------------------------
# Source utilities
# ---------------------------------------------------------------------------

ANSI_INSIDE_STRING_RE = re.compile(r'\\033\[[0-9;]*m')


def _strip_ansi_literals(code: str) -> str:
    """Remove literal "\\033[...m" sequences from code (for readability)."""
    return ANSI_INSIDE_STRING_RE.sub("", code)


def read_source(rel_path: str, start: int | None = None, end: int | None = None) -> str:
    text = (ROOT / rel_path).read_text(encoding="utf-8")
    lines = text.splitlines()
    if start is not None and end is not None:
        lines = lines[start - 1 : end]
    return _strip_ansi_literals("\n".join(lines))


def add_function_block(doc: Document, heading: str, description: str,
                       code_path: str, code_range: tuple[int, int],
                       diagram: str | None = None, diagram_caption: str | None = None):
    doc.add_heading(heading, level=3)
    add_para(doc, description)
    if diagram:
        add_figure(doc, ASSETS / diagram, diagram_caption or heading)
    add_code(doc, read_source(code_path, *code_range))


def add_test_case(doc: Document, index: int, description: str, screenshot: str, caption: str,
                  width_inches: float = 5.5):
    h = doc.add_heading(f"Тест {index}", level=3)
    h.paragraph_format.keep_with_next = True
    p = add_para(doc, description)
    p.paragraph_format.keep_with_next = True
    add_figure(doc, ASSETS / screenshot, caption, width_inches=width_inches)


# ---------------------------------------------------------------------------
# Sections
# ---------------------------------------------------------------------------

def section1(doc: Document):
    doc.add_heading("1. ЦЕЛОЧИСЛЕННАЯ АРИФМЕТИКА И ОДНОМЕРНЫЙ МАССИВ", level=1)

    doc.add_heading("1.1. ПОСТАНОВКА ЗАДАНИЯ", level=2)
    add_para(
        doc,
        "Условие. В заданной последовательности натуральных чисел размера n найти "
        "наибольшее число, для которого сумма минимальной и максимальной цифр его "
        "десятичной записи является чётным числом. Массив хранить в динамически "
        "выделяемой памяти; работать с числами без преобразования в строку и без "
        "дополнительного массива цифр.",
    )
    add_para(doc, "Ограничения и требования к программе:", bold=True, indent=False)
    add_bullets(doc, [
        "использовать только учебный материал лекций;",
        "не создавать дополнительный массив для хранения цифр числа;",
        "не задавать число как символьный массив, не использовать объекты класса string;",
        "все повторяющиеся действия оформить в виде функций;",
        "хранить последовательность в динамическом массиве int*;",
        "результат выводить в доступном и понятном виде.",
    ])

    doc.add_heading("1.2. АЛГОРИТМ И ПРОГРАММА", level=2)
    add_para(
        doc,
        "Программа реализована в классе NumberArray. Конструктор запрашивает размер, "
        "выделяет массив через new int[size] и вызывает inputArray. Метод findAnswer "
        "выполняет линейный поиск: для каждого элемента вычисляются минимальная и "
        "максимальная цифры и проверяется чётность их суммы. Среди подходящих чисел "
        "выбирается наибольшее. Деструктор освобождает память через delete[].",
    )

    doc.add_heading("1.2.1. Класс NumberArray и работа с памятью", level=3)
    add_para(doc,
        "Заголовочный файл описывает интерфейс класса: приватные поля (указатель на "
        "массив и его размер) и публичные методы конструктора, деструктора и запуска "
        "алгоритма."
    )
    add_code(doc, read_source("integerArithmetic/NumberArray.h", 3, 27))
    add_para(doc,
        "Конструктор и деструктор отвечают за жизненный цикл динамического массива:"
    )
    add_code(doc, read_source("integerArithmetic/NumberArray.cpp", 5, 18))

    doc.add_heading("1.2.2. Функции inputArray и printArray", level=3)
    add_para(doc,
        "Ввод и вывод элементов выполняются через указатель ptr, который перемещается "
        "от начала массива до array + size. Это демонстрирует прямую работу с "
        "адресной арифметикой."
    )
    add_code(doc, read_source("integerArithmetic/NumberArray.cpp", 20, 45))

    doc.add_heading("1.2.3. Функции getMinDigit и getMaxDigit", level=3)
    add_para(doc,
        "Обе функции извлекают цифры целочисленным делением на 10 и сравнивают их с "
        "текущим экстремумом. Начальное значение minDigit равно 9, maxDigit — 0, что "
        "гарантирует обновление при первом же шаге для положительных чисел."
    )
    add_figure(doc, ASSETS / "lab1_getMinDigit.png",
               "Блок-схема функции getMinDigit", width_inches=4.0)
    add_code(doc, read_source("integerArithmetic/NumberArray.cpp", 46, 78))

    doc.add_heading("1.2.4. Функция isSuitable", level=3)
    add_para(doc,
        "Проверяет ключевое условие задачи: число является подходящим, если сумма его "
        "минимальной и максимальной цифр — чётная."
    )
    add_figure(doc, ASSETS / "lab1_isSuitable.png",
               "Блок-схема функции isSuitable", width_inches=4.5)
    add_code(doc, read_source("integerArithmetic/NumberArray.cpp", 80, 83))

    doc.add_heading("1.2.5. Функция findAnswer", level=3)
    add_para(doc,
        "Главная функция обработки массива. Поддерживает флаг found и текущий "
        "максимум answer; возвращает -1, если подходящих чисел нет, и максимальное "
        "подходящее число в противном случае."
    )
    add_figure(doc, ASSETS / "lab1_findAnswer.png",
               "Блок-схема функции findAnswer")
    add_code(doc, read_source("integerArithmetic/NumberArray.cpp", 85, 110))

    doc.add_heading("1.2.6. Функция run", level=3)
    add_para(doc,
        "Точка входа алгоритма: выводит массив, вычисляет ответ и печатает результат "
        "с цветовым оформлением через ANSI escape-последовательности."
    )
    add_code(doc, read_source("integerArithmetic/NumberArray.cpp", 112, 137))

    doc.add_heading("1.3. ТЕСТИРОВАНИЕ ПРОЕКТА", level=2)
    add_table(doc,
        ["Тест", "Входные данные", "Ожидаемый результат"],
        [
            ["1", "5 элементов: 12 45 88 23 71", "Largest suitable number = 88"],
            ["2", "3 элемента: 13 17 19", "Suitable numbers not found."],
            ["3", "1 элемент: 246", "Largest suitable number = 246"],
        ],
        "Тесты лабораторной работы 1",
    )

    add_test_case(doc, 1,
                  "Массив из пяти чисел. Для 88 minDigit=8, maxDigit=8, сумма 16 — "
                  "чётная; для 12: 1+2=3 — нечётная. Наибольшее подходящее число — 88.",
                  "lab1_test1.png",
                  "Вывод программы для теста 1 — найдено максимальное подходящее число")
    add_test_case(doc, 2,
                  "Ни одно из чисел 13, 17, 19 не имеет чётной суммы min и max цифр. "
                  "Программа сообщает об отсутствии подходящих значений.",
                  "lab1_test2.png",
                  "Вывод программы для теста 2 — подходящие числа не найдены")
    add_test_case(doc, 3,
                  "Массив из одного числа 246: min=2, max=6, сумма 8 — чётная. Ответом "
                  "становится сам этот элемент.",
                  "lab1_test3.png",
                  "Вывод программы для теста 3 — единственное подходящее число")


def section2(doc: Document):
    doc.add_heading("2. ДВУМЕРНЫЕ МАССИВЫ", level=1)

    doc.add_heading("2.1. ПОСТАНОВКА ЗАДАНИЯ", level=2)
    add_para(doc,
        "Задание 5 (сложность B). Дана прямоугольная целочисленная матрица размера "
        "(n × m). Требуется вставить строку из нулей после последней убывающей строки "
        "матрицы, содержащей максимальное количество элементов, у которых минимальная "
        "цифра — чётная. При равенстве количества выбирается последняя из подходящих "
        "убывающих строк."
    )
    add_para(doc, "Ограничения и требования к программе:", bold=True, indent=False)
    add_bullets(doc, [
        "использовать статический массив data[MAX_ROWS][MAX_COLS];",
        "доступ к строкам выполнять через массив указателей rowPtrs[MAX_ROWS];",
        "при вставке строки перемещать только адреса, дополнительный массив не создавать;",
        "предусмотреть загрузку из файла matrix.txt и меню для запуска задания;",
        "все повторяющиеся действия оформить в функции.",
    ])

    doc.add_heading("2.2. АЛГОРИТМ И ПРОГРАММА", level=2)

    doc.add_heading("2.2.1. Хранение матрицы", level=3)
    add_para(doc,
        "Двумерный массив data хранит значения последовательно, а массив rowPtrs "
        "содержит указатели на начало каждой строки. Работа со строками (сравнение, "
        "вставка, перестановка) сводится к манипуляциям указателями, что позволяет "
        "не копировать данные."
    )
    add_figure(doc, ASSETS / "lab2_memory.png",
               "Организация матрицы: data[][] и rowPtrs[]", width_inches=6.0)
    add_code(doc, read_source("lab2/Matrix.h", 7, 15))

    doc.add_heading("2.2.2. Функции minDigit и isDescending", level=3)
    add_para(doc,
        "Функция minDigit возвращает минимальную десятичную цифру числа (учитывая знак). "
        "Функция isDescending проверяет невозрастание элементов строки слева направо."
    )
    add_code(doc, read_source("lab2/Matrix.cpp", 725, 768))

    doc.add_heading("2.2.3. Функция countEvenMinDigit", level=3)
    add_para(doc,
        "Считает количество элементов строки, у которых минимальная цифра является "
        "чётной. Используется указатель p, пробегающий диапазон [row, row + cols)."
    )
    add_code(doc, read_source("lab2/Matrix.cpp", 770, 785))

    doc.add_heading("2.2.4. Функция findTargetRow", level=3)
    add_para(doc,
        "Основной поиск: перебирает все строки через rowPtrs, отбирает убывающие с "
        "count > 0 и выбирает последнюю с максимальным count. При отсутствии таких "
        "строк возвращает nullptr."
    )
    add_figure(doc, ASSETS / "lab2_findTargetRow.png",
               "Блок-схема функции findTargetRow")
    add_code(doc, read_source("lab2/Matrix.cpp", 787, 815))

    doc.add_heading("2.2.5. Функция insertZeroRowAfter", level=3)
    add_para(doc,
        "Обнуляет свободную ячейку data[rowCount], а затем сдвигает указатели rowPtrs "
        "вправо, вставляя адрес новой строки сразу после найденной target. "
        "Дополнительный массив не создаётся."
    )
    add_figure(doc, ASSETS / "lab2_insertZeroRow.png",
               "Блок-схема функции insertZeroRowAfter")
    add_code(doc, read_source("lab2/Matrix.cpp", 1140, 1169))

    doc.add_heading("2.2.6. Функция insertZeroRowTask", level=3)
    add_para(doc,
        "Объединяет всё в единый сценарий: выводит матрицу до обработки, ищет "
        "target-строку, вставляет нулевую строку и печатает матрицу после."
    )
    add_code(doc, read_source("lab2/Matrix.cpp", 1171, 1205))

    doc.add_heading("2.3. ТЕСТИРОВАНИЕ ПРОЕКТА", level=2)
    add_table(doc,
        ["Файл", "Сценарий", "Ожидание"],
        [
            ["01_odna_podhodyashaya.txt", "Одна убывающая строка с count > 0", "Нулевая строка после строки 1"],
            ["02_poslednyaya_iz_ravnyh.txt", "Несколько убывающих с равным count", "Выбирается последняя"],
            ["03_net_ubyvayushih.txt", "Убывающих строк нет", "Вставка не выполняется"],
            ["04_net_podhodyashih.txt", "count = 0 у всех убывающих", "Вставка не выполняется"],
        ],
        "Автоматические тесты лабораторной работы 2",
        col_widths=[6.5, 5.0, 4.5],
    )
    add_test_case(doc, 1,
                  "Базовый случай: единственная убывающая строка с count > 0. Нулевая "
                  "строка вставляется сразу после неё.",
                  "lab2_test1.png",
                  "Вывод теста 1 (файл 01_odna_podhodyashaya.txt)")
    add_test_case(doc, 2,
                  "Несколько убывающих строк имеют одинаковый максимальный count. Согласно "
                  "условию задачи выбирается последняя.",
                  "lab2_test2.png",
                  "Вывод теста 2 (файл 02_poslednyaya_iz_ravnyh.txt)")
    add_test_case(doc, 3,
                  "Ни одна из строк не является убывающей, поэтому вставка не производится.",
                  "lab2_test3.png",
                  "Вывод теста 3 (файл 03_net_ubyvayushih.txt)")
    add_test_case(doc, 4,
                  "Убывающие строки есть, но у всех минимальные цифры нечётные, поэтому "
                  "count = 0. Вставка также не выполняется.",
                  "lab2_test4.png",
                  "Вывод теста 4 (файл 04_net_podhodyashih.txt)")


def section3(doc: Document):
    doc.add_heading("3. СИМВОЛЬНЫЕ МАССИВЫ", level=1)

    doc.add_heading("3.1. ПОСТАНОВКА ЗАДАНИЯ", level=2)
    add_para(doc,
        "Вариант 5. Дан произвольный текст на английском языке. Требуется выделить "
        "все английские слова в динамический массив указателей char** и отсортировать "
        "их по возрастанию по составному ключу: сначала по числу гласных (a, e, i, "
        "o, u, y), затем по первой гласной, затем лексикографически (strcmp). "
        "Сортировка должна выполняться перестановкой указателей, без копирования "
        "содержимого строк."
    )
    add_para(doc, "Ограничения и требования к программе:", bold=True, indent=False)
    add_bullets(doc, [
        "не использовать объекты класса std::string;",
        "каждое слово хранить в отдельной динамической строке char[];",
        "массив слов представлять как динамический char**;",
        "применять эффективную сортировку слиянием;",
        "обмен слов реализовать через перестановку указателей.",
    ])

    doc.add_heading("3.2. АЛГОРИТМ И ПРОГРАММА", level=2)

    doc.add_heading("3.2.1. Хранение текста и массива слов", level=3)
    add_para(doc,
        "Класс WordArray хранит указатель на копию текста (char* text) и динамический "
        "массив указателей char** words. Каждое слово — отдельная динамическая строка."
    )
    add_figure(doc, ASSETS / "lab3_memory.png",
               "Организация памяти: char* text и char** words", width_inches=6.0)
    add_code(doc, read_source("lab3/WordArray.h", 3, 48))

    doc.add_heading("3.2.2. Функции isVowel, countVowels, firstVowel", level=3)
    add_para(doc,
        "Функции классифицируют символы и вычисляют характеристики слов, используемые "
        "при сравнении: количество гласных и первая гласная."
    )
    add_code(doc, read_source("lab3/WordArray.cpp", 108, 137))

    doc.add_heading("3.2.3. Функция compareWords", level=3)
    add_para(doc,
        "Функция сравнения задаёт порядок слов при сортировке по трёхуровневому ключу: "
        "число гласных, первая гласная, strcmp."
    )
    add_figure(doc, ASSETS / "lab3_compareWords.png",
               "Блок-схема функции compareWords")
    add_code(doc, read_source("lab3/WordArray.cpp", 138, 153))

    doc.add_heading("3.2.4. Функция extractWords", level=3)
    add_para(doc,
        "Сканирует text, определяет границы английских слов и сохраняет их указатели "
        "в динамическом массиве words. Ёмкость массива увеличивается вдвое при "
        "переполнении (стратегия удвоения)."
    )
    add_code(doc, read_source("lab3/WordArray.cpp", 280, 318))

    doc.add_heading("3.2.5. Функции mergeRanges и mergeSortRange", level=3)
    add_para(doc,
        "Реализуют сортировку слиянием. При слиянии временные буферы содержат только "
        "адреса строк, поэтому копирования данных не происходит."
    )
    add_figure(doc, ASSETS / "lab3_mergeSort.png",
               "Блок-схема функции mergeSortRange")
    add_code(doc, read_source("lab3/WordArray.cpp", 320, 385))

    doc.add_heading("3.2.6. Функция sortWords и processTask", level=3)
    add_para(doc,
        "sortWords запускает mergeSortRange по всему массиву. processTask объединяет "
        "весь пайплайн: подсветка текста, выделение слов, сортировка и вывод."
    )
    add_code(doc, read_source("lab3/WordArray.cpp", 376, 438))

    doc.add_heading("3.3. ТЕСТИРОВАНИЕ ПРОЕКТА", level=2)
    add_table(doc,
        ["Тест", "Вход", "Ожидаемый порядок"],
        [
            ["1", "hi education rhythm", "hi, rhythm, education"],
            ["2", "dog ant fun is", "ant, is, dog, fun"],
            ["3", "apple amber", "amber, apple"],
            ["4", "myth syzygy rhythm", "myth, rhythm, syzygy"],
            ["5", "Hello, world! Programming IS fun.", "is, world, fun, hello, programming"],
        ],
        "Автоматические тесты лабораторной работы 3",
        col_widths=[1.6, 6.5, 6.5],
    )
    add_test_case(doc, 1,
                  "Первый ключ — count гласных: у hi(1) и rhythm(1) их меньше, чем у "
                  "education(5), поэтому они идут раньше.",
                  "lab3_test1.png",
                  "Вывод теста 1 — сортировка по числу гласных")
    add_test_case(doc, 2,
                  "При равном count=1 порядок определяется первой гласной: a < i < o < u.",
                  "lab3_test2.png",
                  "Вывод теста 2 — сортировка по первой гласной")
    add_test_case(doc, 3,
                  "У amber и apple совпадают count и первая гласная. Порядок задаёт "
                  "strcmp: 'amber' < 'apple'.",
                  "lab3_test3.png",
                  "Вывод теста 3 — тай-брейк по strcmp")
    add_test_case(doc, 4,
                  "Буква y считается гласной. У syzygy три гласные (y, y, y), поэтому "
                  "она идёт после myth и rhythm (у которых по одной y).",
                  "lab3_test4.png",
                  "Вывод теста 4 — учёт буквы y как гласной")
    add_test_case(doc, 5,
                  "Многострочный текст со знаками препинания и заглавными буквами: "
                  "программа корректно выделяет и приводит слова к нижнему регистру.",
                  "lab3_test5.png",
                  "Вывод теста 5 — комплексный текст")


def section4(doc: Document):
    doc.add_heading("4. РЕКУРСИЯ", level=1)

    doc.add_heading("4.1. ПОСТАНОВКА ЗАДАНИЯ", level=2)
    add_para(doc,
        "Фишка стоит в начале целочисленной прямой длины N. За один ход она может "
        "переместиться вперёд на 1, 2, …, K позиций. Требуется определить число "
        "различных способов дойти до позиции N, реализовав как рекурсивный, так и "
        "нерекурсивный (методом динамического программирования) алгоритмы, а также "
        "вывести все возможные комбинации ходов."
    )
    add_para(doc, "Ограничения и требования к программе:", bold=True, indent=False)
    add_bullets(doc, [
        "реализовать оба алгоритма — рекурсивный и итеративный;",
        "предусмотреть валидацию входных данных (n ≥ 0, k ≥ 1);",
        "вывести число путей и все комбинации ходов;",
        "сравнить результаты рекурсивного и итеративного алгоритмов;",
        "покрыть решение набором автотестов.",
    ])

    doc.add_heading("4.2. АЛГОРИТМ И ПРОГРАММА", level=2)

    doc.add_heading("4.2.1. Функция validateInput", level=3)
    add_para(doc,
        "Проверяет корректность аргументов и выбрасывает std::invalid_argument при "
        "нарушении инвариантов."
    )
    add_code(doc, read_source("lab4/PathCounter.cpp", 11, 22))

    doc.add_heading("4.2.2. Функция countPathsRecursive", level=3)
    add_para(doc,
        "Рекурсивная реализация: базовый случай n = 0 даёт один пустой путь; для "
        "n > 0 суммируется число путей для всех возможных шагов от 1 до min(k, n)."
    )
    add_figure(doc, ASSETS / "lab4_recursive.png",
               "Блок-схема функции countPathsRecursive")
    add_code(doc, read_source("lab4/PathCounter.cpp", 42, 58))

    doc.add_heading("4.2.3. Функция countPathsIterative", level=3)
    add_para(doc,
        "Динамическое программирование снизу вверх: массив dp[i] хранит число путей "
        "до позиции i. Ответ — dp[n]. Сложность O(n·k) по времени и O(n) по памяти."
    )
    add_figure(doc, ASSETS / "lab4_dp.png",
               "Блок-схема функции countPathsIterative")
    add_code(doc, read_source("lab4/PathCounter.cpp", 60, 76))

    doc.add_heading("4.2.4. Функции collectPathsRecursive и findAllPaths", level=3)
    add_para(doc,
        "Рекурсивный перебор с возвратом (backtracking) собирает все комбинации ходов. "
        "Текущая комбинация хранится в векторе current; при достижении финиша копия "
        "добавляется в общий результат all."
    )
    add_code(doc, read_source("lab4/PathCounter.cpp", 24, 39))
    add_code(doc, read_source("lab4/PathCounter.cpp", 78, 86))

    doc.add_heading("4.2.5. Функция printPaths", level=3)
    add_para(doc,
        "Форматированный вывод всех комбинаций. Для наглядности сумма шагов "
        "подсвечивается зелёным, если совпадает с целевой позицией."
    )
    add_code(doc, read_source("lab4/PathCounter.cpp", 88, 131))

    doc.add_heading("4.2.6. Сравнение алгоритмов", level=3)
    add_table(doc,
        ["Алгоритм", "Время", "Память", "Особенности"],
        [
            ["Рекурсивный", "O(K^N)", "O(N) — глубина стека", "Прямая реализация условия задачи"],
            ["DP (итеративный)", "O(N·K)", "O(N) — массив dp", "Значительно эффективнее"],
            ["Перебор путей", "O(P·N)", "O(N) — вектор current", "P — число комбинаций"],
        ],
        "Сравнение алгоритмов подсчёта путей",
        col_widths=[3.5, 3.0, 3.5, 6.5],
    )
    add_figure(doc, ASSETS / "lab4_paths_chart.png",
               "Зависимость числа путей от N при разных K", width_inches=6.0)

    doc.add_heading("4.3. ТЕСТИРОВАНИЕ ПРОЕКТА", level=2)
    add_table(doc,
        ["Тест", "N", "K", "Ожидание", "Комментарий"],
        [
            ["1", "3", "2", "3", "Пример из условия: (1,1,1), (2,1), (1,2)"],
            ["2", "0", "2", "1", "Фишка уже на финише — один пустой путь"],
            ["3", "5", "1", "1", "Единственная последовательность 1+1+1+1+1"],
            ["4", "4", "4", "8", "Общий случай: восемь комбинаций"],
            ["5", "5", "3", "13", "Проверка совпадения DP и рекурсии"],
        ],
        "Автоматические тесты лабораторной работы 4",
        col_widths=[1.3, 1.3, 1.3, 2.0, 10.0],
    )
    add_test_case(doc, 1,
                  "Сводный вывод пяти рекурсивных тестов. Для каждого теста программа "
                  "сообщает ожидаемое и полученное число путей и печатает все комбинации "
                  "ходов (в сводке показаны первые тесты).",
                  "lab4_recursive_summary.png",
                  "Вывод рекурсивных тестов (сводка)",
                  width_inches=4.2)
    add_test_case(doc, 2,
                  "Аналогичная сводка для нерекурсивного (DP) алгоритма. Ответы всех "
                  "тестов совпадают с рекурсивным решением, что подтверждает корректность "
                  "обеих реализаций.",
                  "lab4_iterative_summary.png",
                  "Вывод нерекурсивных тестов (сводка)",
                  width_inches=4.2)


def section5(doc: Document):
    doc.add_heading("5. ФАЙЛЫ И МЕТОД ГАУССА", level=1)

    doc.add_heading("5.1. ПОСТАНОВКА ЗАДАНИЯ", level=2)
    add_para(doc,
        "Разработать программу решения системы линейных алгебраических уравнений вида "
        "Ax = b методом Гаусса с выбором главного элемента по столбцу. Матрица коэффициентов "
        "A размера n × n и столбец свободных членов b размера n должны храниться не в "
        "оперативной памяти, а в отдельных двоичных файлах; все операции прямого и "
        "обратного хода выполняются построчно через файловый ввод-вывод."
    )
    add_para(doc, "Ограничения и требования к программе:", bold=True, indent=False)
    add_bullets(doc, [
        "хранить матрицу и вектор в отдельных бинарных файлах matrix.bin и vector.bin;",
        "обеспечить перестановку строк матрицы и столбца свободных членов на диске;",
        "прямой и обратный ход выполнять построчно, не загружая матрицу целиком;",
        "поддерживать интерактивный ввод системы (TUI-редактор сетки);",
        "результат записывать в текстовый файл result.txt;",
        "предусмотреть набор автотестов для успешных и ошибочных случаев;",
        "максимальный размер системы n = 20, точность сравнения EPS = 10^-9.",
    ])

    doc.add_heading("5.2. АЛГОРИТМ И ПРОГРАММА", level=2)

    doc.add_heading("5.2.1. Организация бинарных файлов", level=3)
    add_para(doc,
        "Каждый бинарный файл начинается со служебной записи int n (размер системы). "
        "За ней следуют n записей: в matrix.bin — строки матрицы по n double, в vector.bin — "
        "n элементов double. Такая структура позволяет обращаться к любой строке или "
        "элементу через seekg/seekp без чтения всего файла."
    )
    add_figure(doc, ASSETS / "lab5_binary_layout.png",
               "Организация бинарных файлов matrix.bin, vector.bin и текстового result.txt",
               width_inches=6.0)
    add_code(doc, read_source("lab5/LinearSystem.h", 7, 38))

    doc.add_heading("5.2.2. Функции доступа к строкам и элементам", level=3)
    add_para(doc,
        "readMatrixRow и writeMatrixRow позволяют читать и записывать одну строку целиком, "
        "readMatrixElement — прочитать один элемент по индексам (i, j). Все три функции "
        "используют seekg/seekp, вычисляя смещение через matrixRowOffset."
    )
    add_code(doc, read_source("lab5/LinearSystem.cpp", 113, 152))

    doc.add_heading("5.2.3. Функция swapMatrixRows", level=3)
    add_para(doc,
        "Меняет местами две строки матрицы прямо на диске. Читает обе строки во временные "
        "буферы rowA/rowB, затем записывает их на позиции друг друга. Аналогичная "
        "swapVectorTerms синхронно переставляет свободные члены."
    )
    add_figure(doc, ASSETS / "lab5_swapRows.png",
               "Блок-схема функции swapMatrixRows", width_inches=4.5)
    add_code(doc, read_source("lab5/LinearSystem.cpp", 154, 170))

    doc.add_heading("5.2.4. Функция solveSystemGauss", level=3)
    add_para(doc,
        "Основной алгоритм. Читает размеры из файлов, проверяет их совпадение, выполняет "
        "прямой ход с выбором главного элемента: для каждого столбца k ищется строка с "
        "максимальным по модулю коэффициентом, при необходимости выполняется перестановка, "
        "затем строки ниже приводятся к треугольному виду записью обратно в файл. Обратный "
        "ход вычисляет корни x[i]; результат записывается в текстовый файл."
    )
    add_figure(doc, ASSETS / "lab5_solveGauss.png",
               "Блок-схема функции solveSystemGauss")
    add_code(doc, read_source("lab5/LinearSystem.cpp", 344, 400))
    add_para(doc, "Продолжение (прямой ход, обратный ход, запись результата):")
    add_code(doc, read_source("lab5/LinearSystem.cpp", 427, 516))

    doc.add_heading("5.2.5. Функция createBinaryFiles", level=3)
    add_para(doc,
        "Интерактивно запрашивает размер системы, запускает TUI-редактор сетки коэффициентов "
        "(inputSystemFromKeyboard), а затем сохраняет матрицу и вектор в бинарные файлы."
    )
    add_code(doc, read_source("lab5/LinearSystem.cpp", 259, 288))

    doc.add_heading("5.3. ТЕСТИРОВАНИЕ ПРОЕКТА", level=2)
    add_table(doc,
        ["Тест", "Файл описания", "Тип", "Ожидание"],
        [
            ["1", "01_success_3x3.txt", "Успех", "A[3×3], b[3] → x = (2, 3, -1)"],
            ["2", "02_pivot_swap.txt", "Успех", "A с нулём на диагонали → перестановка строк"],
            ["3", "03_bad_sizes.txt", "Ошибка", "Размеры матрицы и вектора не совпадают"],
            ["4", "04_singular.txt", "Ошибка", "Вырожденная система (пропорциональные строки)"],
            ["5", "05_inconsistent.txt", "Ошибка", "Несовместная система (параллельные уравнения)"],
        ],
        "Автоматические тесты лабораторной работы 5",
        col_widths=[1.5, 4.5, 2.5, 7.5],
    )
    add_test_case(doc, 1,
                  "Эталонная система 3×3 с очевидным решением. Программа выполняет прямой ход "
                  "с выбором главного элемента и обратный ход; корни совпадают с ожиданием.",
                  "lab5_test1.png",
                  "Вывод автотеста 1 — успешное решение системы 3×3")
    add_test_case(doc, 2,
                  "Первый коэффициент a[0][0] равен нулю, поэтому алгоритм переставляет строки "
                  "и выбирает вторую строку как опорную. Проверяется корректность pivoting.",
                  "lab5_test2.png",
                  "Вывод автотеста 2 — перестановка строк (pivoting)")
    add_test_case(doc, 3,
                  "Матрица 2×2 и вектор из 3 элементов. Программа обнаруживает несовпадение "
                  "размеров и сообщает об ошибке до попытки решения.",
                  "lab5_test3.png",
                  "Вывод автотеста 3 — несовпадение размеров")
    add_test_case(doc, 4,
                  "Вторая строка матрицы пропорциональна первой. При прямом ходе главный элемент "
                  "оказывается нулевым, и программа сообщает, что систему решить нельзя.",
                  "lab5_test4.png",
                  "Вывод автотеста 4 — вырожденная система")
    add_test_case(doc, 5,
                  "Уравнения параллельны, но правые части различны. Как и в тесте 4, обнаруживается "
                  "нулевой главный элемент, решение невозможно.",
                  "lab5_test5.png",
                  "Вывод автотеста 5 — несовместная система")


def section6(doc: Document):
    doc.add_heading("6. МЕТОД ТРАПЕЦИЙ И МЕТОД НЬЮТОНА", level=1)

    doc.add_heading("6.1. ПОСТАНОВКА ЗАДАНИЯ", level=2)
    add_para(doc,
        "Вариант 5, повышенная сложность (задание 3.5.2). Требуется вычислить определённый "
        "интеграл I(c) = ∫[1;2] √(x² + c²) / x dx для набора значений параметра c "
        "составной формулой трапеций с автоматическим выбором числа разбиений. Для "
        "вычисления квадратного корня в подынтегральной функции запрещено использовать "
        "std::sqrt: значение √t находится как корень уравнения y² − t = 0 методом Ньютона. "
        "Численный интеграл сверяется с аналитическим значением."
    )
    add_para(doc, "Ограничения и требования к программе:", bold=True, indent=False)
    add_bullets(doc, [
        "составную формулу трапеций реализовать для произвольной f, передаваемой указателем на функцию;",
        "√t вычислять только методом Ньютона по формуле y_(k+1) = y_k − (y_k² − t)/(2·y_k);",
        "критерий остановки метода Ньютона: |y_(k+1) − y_k| < eps_sqrt;",
        "отделение корня уравнения y² − t = 0 проверять явно (isRootBracketed);",
        "адаптивно удваивать число разбиений n → 2n → 4n → … до |I_prev − I_curr| < eps_integral;",
        "сравнивать численный результат с аналитическим значением интеграла;",
        "покрыть решение автотестами (трапеции, Ньютон, адаптив, вариант 5).",
    ])

    doc.add_heading("6.2. АЛГОРИТМ И ПРОГРАММА", level=2)

    doc.add_heading("6.2.1. Метод Ньютона (newtonSolve)", level=3)
    add_para(doc,
        "Универсальная реализация метода Ньютона для решения уравнения f(y) = 0. Функция "
        "принимает указатели на f и её производную df, начальное приближение y0 и точность "
        "eps. Останавливается при |y_(k+1) − y_k| < eps либо при достижении лимита итераций. "
        "Дополнительно защищена от деления на ноль при малой производной."
    )
    add_figure(doc, ASSETS / "lab6_newton.png",
               "Блок-схема универсального метода Ньютона")
    add_code(doc, read_source("lab6/main.cpp", 80, 107))

    doc.add_heading("6.2.2. Функции f_sqrt и df_sqrt", level=3)
    add_para(doc,
        "Уравнение y² − t = 0 представлено в виде функции f_sqrt(y) = y² − t с производной "
        "df_sqrt(y) = 2·y. Значение t передаётся через глобальную переменную t_value — так "
        "сохраняется требование передавать функции только по указателю без замыкания параметров."
    )
    add_code(doc, read_source("lab6/main.cpp", 54, 65))

    doc.add_heading("6.2.3. Функция sqrtByNewton", level=3)
    add_para(doc,
        "Оболочка над newtonSolve для вычисления √t. Проверяет корректность аргумента "
        "(t ≥ 0), отделение корня на [0; t], а затем запускает метод Ньютона с начальным "
        "приближением y0 = t. Если t = 0 — возвращает 0 сразу."
    )
    add_figure(doc, ASSETS / "lab6_sqrtNewton.png",
               "Блок-схема функции sqrtByNewton", width_inches=5.5)
    add_code(doc, read_source("lab6/main.cpp", 109, 127))

    doc.add_heading("6.2.4. Составная формула трапеций (compositeTrapezoid)", level=3)
    add_para(doc,
        "Реализация составной формулы трапеций для произвольной подынтегральной функции: "
        "I ≈ h · (f(a)/2 + f(a+h) + … + f(b−h) + f(b)/2), где h = (b − a) / n. Функция "
        "принимает f через указатель, что позволяет использовать её как для тестовой f(x) = x, "
        "так и для подынтегральной функции варианта 5."
    )
    add_figure(doc, ASSETS / "lab6_trapezoid.png",
               "Блок-схема составной формулы трапеций", width_inches=5.0)
    add_code(doc, read_source("lab6/main.cpp", 163, 183))

    doc.add_heading("6.2.5. Адаптивный метод трапеций (adaptiveTrapezoid)", level=3)
    add_para(doc,
        "Повышенная сложность задания. Вычисляет интеграл при n, 2n, 4n, … разбиениях и "
        "останавливается, когда разность соседних приближений становится меньше eps_integral. "
        "Возвращает результат вместе со служебной информацией: начальное и итоговое число "
        "разбиений, последняя разность, флаг сходимости."
    )
    add_figure(doc, ASSETS / "lab6_adaptive.png",
               "Блок-схема адаптивного метода трапеций")
    add_code(doc, read_source("lab6/main.cpp", 185, 240))

    doc.add_heading("6.2.6. Подынтегральная функция варианта 5", level=3)
    add_para(doc,
        "Функция integrand возвращает значение √(x² + c²) / x, причём корень вычисляется "
        "через sqrtByNewton, а параметр c берётся из глобальной переменной c_value. Для "
        "аналитической проверки используется первообразная "
        "F(x) = √(x² + c²) + (c/2)·ln((√(x² + c²) − c)/(√(x² + c²) + c))."
    )
    add_figure(doc, ASSETS / "lab6_variant5_chart.png",
               "Подынтегральная функция варианта 5 для нескольких значений c", width_inches=6.0)
    add_code(doc, read_source("lab6/main.cpp", 135, 161))

    doc.add_heading("6.2.7. Сценарий runVariant5", level=3)
    add_para(doc,
        "Основной сценарий: для каждого c из диапазона 1.0..3.0 с шагом 0.2 запускается "
        "adaptiveTrapezoid, вычисляется аналитическое значение, ошибка |I − I_an| и последняя "
        "разность |dI|. Результаты выводятся таблицей, что позволяет визуально оценить "
        "сходимость метода."
    )
    add_code(doc, read_source("lab6/main.cpp", 571, 606))

    doc.add_heading("6.3. ТЕСТИРОВАНИЕ ПРОЕКТА", level=2)
    add_table(doc,
        ["Тест", "Проверяет", "Вход", "Ожидание"],
        [
            ["1", "compositeTrapezoid", "f(x) = x, a=0, b=1, n=10", "I = 0.5, |I − 0.5| < 1e-12"],
            ["2", "sqrtByNewton", "t=10, eps=1e-10", "|√10 − 3.16227…| < 1e-9"],
            ["3", "adaptiveTrapezoid", "f(x) = x, n0=2, eps=1e-8", "converged=true, I≈0.5"],
            ["4", "Вариант 5 vs аналитика", "c=2, [1;2], n0=10, eps=1e-5", "|I_num − I_an| < 1e-4"],
        ],
        "Автоматические тесты лабораторной работы 6",
        col_widths=[1.4, 4.5, 5.5, 5.5],
    )
    add_test_case(doc, 1,
                  "Проверка формулы трапеций на линейной функции: точный интеграл ∫[0;1] x dx = 0.5. "
                  "Погрешность в пределах машинного эпсилона — формула реализована корректно.",
                  "lab6_test1.png",
                  "Вывод автотеста 1 — трапеции на f(x) = x", width_inches=5.0)
    add_test_case(doc, 2,
                  "Метод Ньютона решает уравнение y² − 10 = 0 и находит √10. Результат совпадает "
                  "со значением std::sqrt(10) с точностью выше запрошенной eps=1e-10.",
                  "lab6_test2.png",
                  "Вывод автотеста 2 — Ньютон для sqrt(10)", width_inches=5.0)
    add_test_case(doc, 3,
                  "Адаптивное удвоение разбиений на линейной функции: критерий |I_prev − I_curr| < 1e-8 "
                  "выполняется уже после первого удвоения, поскольку трапеции точны на линейной f.",
                  "lab6_test3.png",
                  "Вывод автотеста 3 — адаптив на f(x) = x", width_inches=5.0)
    add_test_case(doc, 4,
                  "Сравнение адаптивного численного решения варианта 5 при c=2 с аналитическим "
                  "значением. Разность |I_num − I_an| ~10⁻⁶, что подтверждает корректность связки "
                  "трапеции + метод Ньютона.",
                  "lab6_test4.png",
                  "Вывод автотеста 4 — вариант 5 vs аналитика", width_inches=5.0)
    add_test_case(doc, 5,
                  "Итоговая таблица результатов основного сценария: для каждого значения c "
                  "показаны число разбиений n_fin (растёт с усложнением подынтегральной функции), "
                  "численный интеграл I_num, аналитический I_an, ошибка |I − I_an| и критерий "
                  "остановки |dI|. Все погрешности порядка 10⁻⁷, что подтверждает работу адаптивного "
                  "метода.",
                  "lab6_variant5_table.png",
                  "Таблица результатов основного сценария (адаптивный метод трапеций)")
    add_test_case(doc, 6,
                  "Отдельная таблица проверки решения уравнения y² − t = 0 для t = 2, 5, 10. "
                  "Остаток err_sqrt = |y² − t| ~10⁻¹⁵ — метод Ньютона находит корень с "
                  "предельной точностью.",
                  "lab6_sqrt_check.png",
                  "Проверка вычисления √t через метод Ньютона")


def section7(doc: Document):
    doc.add_heading("7. ОДНОНАПРАВЛЕННЫЙ СПИСОК СТУДЕНТОВ", level=1)

    doc.add_heading("7.1. ПОСТАНОВКА ЗАДАНИЯ", level=2)
    add_para(doc,
        "Разработать программу управления однонаправленным связным списком записей о "
        "студентах. Каждый элемент списка содержит фамилию, имя и массив из пяти оценок "
        "(целые числа от 1 до 10). Программа должна поддерживать создание списка с "
        "клавиатуры, сортировку по среднему баллу (по возрастанию или убыванию), а также "
        "вставку нового студента с сохранением текущего порядка сортировки."
    )
    add_para(doc, "Ограничения и требования к программе:", bold=True, indent=False)
    add_bullets(doc, [
        "хранить студентов в связном списке (используется std::list<StudentRecord>);",
        "средний балл вычислять как арифметическое среднее пяти оценок;",
        "проверять оценки на принадлежность диапазону [1; 10] (иначе invalid_argument);",
        "сортировка должна быть устойчивой — при равных средних порядок ввода сохраняется;",
        "вставка выполняется через std::find_if с сохранением инварианта отсортированности;",
        "поддерживать оба порядка (возрастание, убывание) через std::function-компаратор;",
        "предусмотреть автотесты для сортировки и вставки в оба порядка.",
    ])

    doc.add_heading("7.2. АЛГОРИТМ И ПРОГРАММА", level=2)

    doc.add_heading("7.2.1. Структура StudentRecord и класс StudentList", level=3)
    add_para(doc,
        "Запись студента реализована как структура с публичными полями фамилии, имени и "
        "std::array<int, 5> для оценок; метод averageGrade возвращает средний балл. "
        "StudentList инкапсулирует std::list<StudentRecord> и предоставляет методы "
        "создания, вывода, сортировки и вставки."
    )
    add_figure(doc, ASSETS / "lab7_list_structure.png",
               "Структура однонаправленного списка StudentList", width_inches=6.5)
    add_code(doc, read_source("lab7/StudentList.h", 11, 44))

    doc.add_heading("7.2.2. Функция averageGrade и валидация оценок", level=3)
    add_para(doc,
        "averageGrade суммирует пять целых оценок и делит на 5.0 — используется как ключ "
        "для сортировки и вставки. validateGrade выбрасывает std::invalid_argument при "
        "оценке вне диапазона [1; 10]."
    )
    add_code(doc, read_source("lab7/StudentList.cpp", 9, 33))

    doc.add_heading("7.2.3. Компараторы среднего балла", level=3)
    add_para(doc,
        "Компаратор задаётся как std::function<bool(double, double)>. Функции "
        "ascendingAverage и descendingAverage возвращают готовые лямбды, "
        "readAverageComparatorFromKeyboard спрашивает у пользователя направление порядка. "
        "comparatorLabel определяет текущее направление по поведению переданного компаратора."
    )
    add_code(doc, read_source("lab7/Comparators.h", 1, 43))

    doc.add_heading("7.2.4. Функция sortByAverage", level=3)
    add_para(doc,
        "Делегирует сортировку методу std::list::sort с лямбдой, вызывающей переданный "
        "компаратор от averageGrade двух записей. std::list::sort реализована как устойчивая "
        "сортировка слиянием со сложностью O(n log n), что подтверждается тестом 3."
    )
    add_figure(doc, ASSETS / "lab7_sort.png",
               "Блок-схема функции sortByAverage", width_inches=4.5)
    add_code(doc, read_source("lab7/StudentList.cpp", 100, 105))

    doc.add_heading("7.2.5. Функция insertSorted", level=3)
    add_para(doc,
        "std::find_if последовательно перебирает элементы и возвращает первый, для которого "
        "compare(newAvg, currentAvg) истинно; в эту позицию через list::insert вставляется "
        "новый студент. Если такого элемента нет, вставка происходит в конец списка."
    )
    add_figure(doc, ASSETS / "lab7_insertSorted.png",
               "Блок-схема функции insertSorted", width_inches=5.0)
    add_code(doc, read_source("lab7/StudentList.cpp", 107, 119))

    doc.add_heading("7.2.6. Функции создания и ввода студента", level=3)
    add_para(doc,
        "createFromKeyboard очищает список, запрашивает количество студентов и по одному "
        "добавляет их через append. readStudentFromKeyboard читает фамилию, имя и пять оценок; "
        "каждая оценка проверяется через validateGrade."
    )
    add_code(doc, read_source("lab7/StudentList.cpp", 50, 68))
    add_code(doc, read_source("lab7/StudentList.cpp", 143, 160))

    doc.add_heading("7.3. ТЕСТИРОВАНИЕ ПРОЕКТА", level=2)
    add_table(doc,
        ["Тест", "Тип", "Что проверяется"],
        [
            ["1", "Сортировка", "Четыре различных дробных средних (5.6, 6.8, 7.6, 8.4)"],
            ["2", "Граничный", "Список из одного элемента (средний 6.8)"],
            ["3", "Стабильность", "Три студента с одинаковым средним 7.2 — сохранение порядка"],
            ["4", "Вставка", "Новый минимум (5.4) и максимум (8.8) в отсортированный список"],
            ["5", "Вставка", "Вставка в середину (7.2) и с равным ключом (7.6)"],
        ],
        "Автоматические тесты лабораторной работы 7",
        col_widths=[1.5, 3.0, 11.5],
    )
    add_test_case(doc, 1,
                  "Четыре студента с дробными средними баллами располагаются в порядке "
                  "возрастания. Проверяется базовая сортировка по ключу.",
                  "lab7_asc_test1.png",
                  "Вывод автотеста 1 (по возрастанию) — сортировка четырёх студентов")
    add_test_case(doc, 3,
                  "Три студента с одинаковым средним баллом 7.2. После сортировки их порядок "
                  "должен совпадать с порядком ввода — это подтверждает устойчивость std::list::sort.",
                  "lab7_asc_test3.png",
                  "Вывод автотеста 3 (по возрастанию) — устойчивость сортировки")
    add_test_case(doc, 4,
                  "В уже отсортированный список из четырёх студентов вставляются новый минимум "
                  "(средний 5.4) и новый максимум (средний 8.8). Проверяется корректность "
                  "insertSorted для граничных позиций.",
                  "lab7_asc_test4.png",
                  "Вывод автотеста 4 (по возрастанию) — вставка на края списка")
    add_test_case(doc, 5,
                  "Вставка в середину списка (средний 7.2) и вставка студента с уже существующим "
                  "средним (7.6). Дубликат ключа добавляется корректно, инвариант сохраняется.",
                  "lab7_asc_test5.png",
                  "Вывод автотеста 5 (по возрастанию) — вставка в середину и при равном ключе")
    add_test_case(doc, 6,
                  "Тот же тест 1, но с компаратором убывания. Результат — обратный порядок средних "
                  "баллов, что подтверждает работу компаратора-стратегии.",
                  "lab7_desc_test1.png",
                  "Вывод автотеста 1 (по убыванию) — сортировка в обратном порядке")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def build_report():
    global FIG_COUNTER, TAB_COUNTER
    FIG_COUNTER = 0
    TAB_COUNTER = 0

    doc = Document()
    setup_sections(doc)
    setup_styles(doc)
    enable_update_fields(doc)

    add_title_page(doc)
    add_toc(doc)
    section1(doc)
    section2(doc)
    section3(doc)
    section4(doc)
    section5(doc)
    section6(doc)
    section7(doc)

    doc.save(OUTPUT)
    print(f"Report saved: {OUTPUT}")
    print(f"Figures: {FIG_COUNTER}, tables: {TAB_COUNTER}")


if __name__ == "__main__":
    build_report()
